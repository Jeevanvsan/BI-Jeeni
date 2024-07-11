import docx
from docx.shared import RGBColor
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import yaml
from PIL import Image


class doc_generator:
    def __init__(self,file):

        with open(f'output/{file}/data_structure.yaml','r') as f:
            self.data = yaml.safe_load(f)
        self.doc = docx.Document() 
        self.file = file

        sec_pr = self.doc.sections[0]._sectPr # get the section properties el
        # create new borders el
        pg_borders = OxmlElement('w:pgBorders')
        # specifies how the relative positioning of the borders should be calculated
        pg_borders.set(qn('w:offsetFrom'), 'page')
        for border_name in ('top', 'left', 'bottom', 'right',): # set all borders
            border_el = OxmlElement(f'w:{border_name}')
            border_el.set(qn('w:val'), 'single') # a single line
            border_el.set(qn('w:sz'), '4') # for meaning of  remaining attrs please look docs
            border_el.set(qn('w:space'), '24')
            border_el.set(qn('w:color'), 'auto')
            pg_borders.append(border_el) # register single border to border el
        sec_pr.append(pg_borders) # apply border changes to section

    def run(self):
        doc = self.doc
        self.front_page()
        doc.add_page_break()
        self.schema()
        doc.add_page_break()
        self.visual_doc()
        doc.save(f'output/{self.file}/{self.file}.docx')

    def front_page(self):
        doc = self.doc
        h1 = doc.add_heading('POWERBI REPORT ANALYSIS',1)
        title_format = h1.runs[0].font
        title_format.size = Pt(54)
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h1.bold = True

        doc.add_picture('img/pbi_logo.png')
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


        file = doc.add_heading(f'\n{self.file}',1)
        title_format = file.runs[0].font
        title_format.size = Pt(26)
        file.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    def schema(self):
        doc = self.doc
        h1 = doc.add_heading('SCHEMA',1)
        title_format = h1.runs[0].font
        title_format.size = Pt(48)
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h1.bold = True

        for tables in self.data['schema']['tables']:
            table_name = tables['table_name']
                
            tbl_n = doc.add_heading('TABLE NAME. : ',3)
            tbl_n.add_run(f'\t{table_name}')
            title_format = tbl_n.runs[0].font
            title_format.size = Pt(22)
            title_format = tbl_n.runs[1].font
            title_format.size = Pt(22)

            table = doc.add_table(rows=1, cols=2)
            # Access the first row to add data
            first_row = table.rows[0]
            first_row.cells[0].text = "COLUMNS"        
            first_row.cells[1].text = "DATA TYPES" 

            # # print(tables['m_code'])

            for col in tables['columns']:
                c_name = col['name']
                c_type = col['datatype']

                row = table.add_row()
                row.cells[0].text = c_name
                row.cells[1].text = c_type

            if "mode" in tables:
                tbl_n = doc.add_heading('MODE : ',3)
                tbl_n.add_run(f'\t{tables["mode"]}')
                title_format = tbl_n.runs[0].font
                title_format.size = Pt(22)
                title_format = tbl_n.runs[1].font
                title_format.size = Pt(22)

            if "m_code" in tables:
                tbl_n = doc.add_heading('M CODE : ',3)
                title_format = tbl_n.runs[0].font
                title_format.size = Pt(22)
                for m in tables["m_code"]:
                    v_type = doc.add_paragraph()
                    v_type.add_run(f'\t{m}') 

            if "DAX" in tables:
                dax = doc.add_heading('DAX : ',3)
                title_format = dax.runs[0].font
                title_format.size = Pt(22)
                for d in tables["DAX"]:
                    # # print(d["expression"])
                    d1 = doc.add_paragraph('Name :',style ='List Bullet 2')
                    d1.add_run(f'\t{d.get("name")}').bold = True

                    d2 = doc.add_paragraph('Expression :',style ='List Bullet 2')
                    if isinstance(d["expression"],list):
                        for e in d["expression"] :
                            v_type = doc.add_paragraph()
                            v_type.add_run(f'\t{e}') 
                    else:
                        v_type = doc.add_paragraph()
                        v_type.add_run(f'\t{d["expression"]}') 

                        
                    

    def visual_doc(self):
        doc = self.doc
        for i,report in enumerate(self.data['layout'],1):
            self.report_name = report["report_name"]
            h1 = doc.add_heading('REPORT NO. : ',1)
            h1.add_run(f'\t{i}')
            title_format = h1.runs[0].font
            title_format.size = Pt(48)
            title_format = h1.runs[1].font
            title_format.size = Pt(48)
            h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            h1.bold = True
            h2 = doc.add_heading('REPORT NAME. : ',3)
            h2.add_run(f'\t{report["report_name"]}')
            title_format = h2.runs[0].font
            title_format.size = Pt(22)
            title_format = h2.runs[1].font
            title_format.size = Pt(22)
            h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            h2.bold = True
            doc.add_page_break()
            self.layout_doc(report['visualContainers'],report)
            doc.add_page_break()
        
        # doc.save('out.docx')
    def add_layout_img(self):
        doc = self.doc
        doc.add_heading('REPORT LAYOUT IMAGE\n', 3).bold = True
        doc.add_picture(f'static/img/layout/{self.file}/{self.report_name}.png')

    def layout_doc(self,report,main_report):
        doc = self.doc
        doc.add_heading('LAYOUT ', 1).bold = True
        self.add_layout_img()


        doc.add_heading('Canvas settings ', 1).bold = True
        for key, value in main_report["canvas_settings"][0].items():
                if "url" in key:
                    try:
                        doc.add_picture(f'output/{self.file}/img/Report/{value}', width=Inches(2))
                    except:
                        pass
                cs = doc.add_paragraph(style ='List Bullet')
                cs.add_run(f'{key} : ').bold = True
                cs.add_run(f'\t{value}')

        if main_report["wallpaper"]:
            doc.add_heading('Wallpaper settings ', 1).bold = True
            for key, value in main_report["wallpaper"][0].items():
                    if "url" in key:
                        try:
                            doc.add_picture(f'output/{self.file}/img/Report/{value}', width=Inches(2))
                        except:
                            pass                  
                    ws = doc.add_paragraph(style ='List Bullet')
                    ws.add_run(f'{key} : ').bold = True
                    ws.add_run(f'\t{value}')
        


        for i,visual in enumerate(report,1):
            for container_ in visual:
                container =  visual[container_]
                title = ''

                if 'title' in container:
                    if container['title'] :
                        title = container['title']
                    else:
                        title = 'default value'
                else:
                    title = container_

                doc.add_heading(f'CONTAINER_{i}', 2) 

                # title 
                t1 = doc.add_paragraph(style ='List Bullet')
                t1.add_run(f'CONTAINER_TITLE : ').bold = True
                t1.add_run(f'\t{title}')
                t1.paragraph_format.right_indent = Inches(0.0)

                # visual type
                v_type = doc.add_paragraph(style ='List Bullet')
                v_type.add_run(f'VISUAL_TYPE : ').bold = True
                v_type.add_run(f'\t{container["visual_type"]}')
                
                # name
                if 'name' in container:
                    v_type = doc.add_paragraph(style ='List Bullet')
                    v_type.add_run(f'Name : ').bold = True
                    v_type.add_run(f'\t{container["name"]}') 

                # author
                if 'author' in container:
                    v_type = doc.add_paragraph(style ='List Bullet')
                    v_type.add_run(f'Author : ').bold = True
                    v_type.add_run(f'\t{container["author"]}') 

                # desc
                if 'desc' in container:
                    v_type = doc.add_paragraph(style ='List Bullet')
                    v_type.add_run(f'Description : ').bold = True
                    v_type.add_run(f'\t{container["desc"]}') 


                # columns
                if 'columns' in container:
                    c_t = doc.add_paragraph(style ='List Bullet')
                    c_t.add_run('COLUMNS ').bold = True
                    if 'card' in container["visual_type"]:
                        cl = doc.add_paragraph(style='List Bullet 2')
                        cl.add_run(f'{container["columns"]}')
                    else:
                        for col in container['columns']:
                            for col_ in col:
                                if col[col_]:
                                    doc.add_paragraph(col_,style='List Bullet 2')
                                    for val in col[col_]:
                                        doc.add_paragraph(val,style='List Bullet 3')
                                    # cl.add_run(f'{col}')
                elif 'text' in container:
                    if container["text"]:
                        c_t = doc.add_paragraph(style ='List Bullet')
                        c_t.add_run('TEXT ').bold = True
                        cl = doc.add_paragraph(style='List Bullet 2')

                        if isinstance(container["text"],list):
                            for txt in container["text"]:
                            # # print(container["text"])
                                cl.add_run(txt)
                    
                        else:
                            cl.add_run(f'{container["text"]}')



                # size
                size = doc.add_paragraph(style ='List Bullet')
                size.add_run('Size ').bold = True
                pos1= container['position']

                # width
                w = doc.add_paragraph(style='List Bullet 2')
                w.add_run('WIDTH : ').bold = True
                w.add_run(f'\t{pos1["width"]}')
                #height
                h = doc.add_paragraph(style='List Bullet 2')
                h.add_run('HEIGHT : ').bold = True
                h.add_run(f'\t{pos1["height"]}')

                # position
                pos = doc.add_paragraph(style ='List Bullet')
                pos.add_run('POSITION ').bold = True
                pos1= container['position']
                #x
                x = doc.add_paragraph(style='List Bullet 2')
                x.add_run('X : ').bold = True
                x.add_run(f'\t{pos1["x"]}')
                #y
                y = doc.add_paragraph(style='List Bullet 2')
                y.add_run('Y : ').bold = True
                y.add_run(f'\t{pos1["y"]}')

                # filters
                # if 'filters' in container:
                #     if container['filters']:
                #         fil = doc.add_paragraph(style ='List Bullet')
                #         fil.add_run('FILTERS ').bold = True
                #         for filter in container['filters']:
                #             ft = doc.add_paragraph('FILTER TYPE : ',style ='List Bullet 2')
                #             ft.add_run(f'\t{filter["filter_type"]}')
                #             ft = doc.add_paragraph('FILTER TYPE : ',style ='List Bullet 2')
                #             ft.add_run(f'\t{filter["filter_type"]}')

                # style
                if 'style' in container:
                    if container['style']:
                        for styles in container['style']:
                            for key, value in styles.items():
                                if value:
                                    styl = doc.add_paragraph(style ='List Bullet')
                                    styl.add_run('STYLE ').bold = True
                                    doc.add_paragraph(key,style ='List Bullet 2')
                                    for val in value:
                                        for k,v in val.items():
                                            v_ = doc.add_paragraph(f'{k} : ',style ='List Bullet 3')
                                            v_.add_run(f'\t{v}')

                if 'image_url' in container:
                    url_ = container["image_url"]
                    v_type = doc.add_paragraph(style ='List Bullet')
                    v_type.add_run(f'Image : ').bold = True
                    v_type.add_run(f'\t{url_}') 
                    doc.add_picture(f'output/{self.file}/img/Report/{url_}',width=Inches(2))
                            # doc.add_picture(f'static/img/report/{url_}',width=Inches(2))x
                ignore_list = ["visual_type","x","y","legend","value","size","secondary_line","X","Y","Category","image_url","position","visual_type","column","columns","text","id"]
                
                for key, value in container.items():
                    if key not in ignore_list:
                        key = key.replace("_"," ")
                        v_type = doc.add_paragraph(style ='List Bullet')
                        v_type.add_run(f'{key} : ').bold = True
                        v_type.add_run(f'\t{value}')
                        
                        
                            

                            

if __name__ == '__main__':
    obj = doc_generator('Zomato-JeevanV_')
    obj.run()


